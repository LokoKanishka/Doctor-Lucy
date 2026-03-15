import os
import sys

try:
    import fitz # PyMuPDF
except ImportError:
    os.system("pip install pymupdf python-docx")
    import fitz

try:
    from docx import Document
except ImportError:
    os.system("pip install python-docx")
    from docx import Document

print("Extracting text from PDF...")
pdf_path = '/home/lucy-ubuntu/Descargas/Obiols - Nuevo curso de lógica y filosofía (1).pdf'
doc = fitz.open(pdf_path)

out_doc = Document()
out_doc.add_heading('Obiols - Nuevo curso de lógica y filosofía', 0)

total_pages = len(doc)
print(f"Total pages: {total_pages}")

for i in range(total_pages):
    page = doc[i]
    text = page.get_text()
    if text.strip():
        out_doc.add_paragraph(text)
    if i % 50 == 0:
        print(f"Processed {i}/{total_pages} pages")

out_path = '/home/lucy-ubuntu/Escritorio/oblois.docx'
out_doc.save(out_path)
print(f"Saved successfully to {out_path}")
